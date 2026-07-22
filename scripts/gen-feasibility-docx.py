#!/usr/bin/env python3
"""Render the feasibility markdown (controlled subset) to .docx via python-docx.

Handled subset: `#`/`##`/`###` headings, `-`/`  -` bullets, `**bold**` inline,
pipe tables, plain paragraphs, `---` rule (skipped). Not a general Markdown parser —
input is repo-controlled (docs/specs/2026-07-06-foundation-model-vision-feasibility.md).

Usage:
  python3 scripts/gen-feasibility-docx.py            # md -> docx
  python3 scripts/gen-feasibility-docx.py --selftest # assert-based render self-check
"""
import re
import sys

from docx import Document

SRC = "/home/rokey/ros2_jazzy_test/docs/specs/2026-07-06-foundation-model-vision-feasibility.md"
OUT = "/home/rokey/ros2_jazzy_test/docs/specs/2026-07-06-foundation-model-vision-feasibility.docx"


def add_runs(par, text):
    """Add `text` to paragraph, toggling bold on ** boundaries (odd segments bold)."""
    for i, seg in enumerate(text.split("**")):
        if seg:
            par.add_run(seg).bold = (i % 2 == 1)


def _is_table_row(line):
    return line.strip().startswith("|")


def _is_separator(cells):
    return all(c and set(c) <= set("-: ") for c in cells)


def render(md_lines, doc):
    """Render markdown lines into `doc`; return (n_headings, n_tables)."""
    n_head = n_table = 0
    i = 0
    while i < len(md_lines):
        raw = md_lines[i].rstrip("\n")
        stripped = raw.strip()
        if not stripped or stripped == "---":
            i += 1
            continue

        if _is_table_row(stripped):  # gather a contiguous table block
            block = []
            while i < len(md_lines) and _is_table_row(md_lines[i]):
                block.append(md_lines[i].strip())
                i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in block]
            rows = [r for r in rows if not _is_separator(r)]
            if rows:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncol)
                tbl.style = "Table Grid"
                for ri, r in enumerate(rows):
                    for ci in range(ncol):
                        p = tbl.rows[ri].cells[ci].paragraphs[0]
                        add_runs(p, r[ci] if ci < len(r) else "")
                        if ri == 0:
                            for run in p.runs:
                                run.bold = True
                n_table += 1
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            doc.add_heading(m.group(2), level=len(m.group(1)))
            n_head += 1
            i += 1
            continue

        if stripped.startswith("- "):
            indent = len(raw) - len(raw.lstrip(" "))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            add_runs(doc.add_paragraph(style=style), stripped[2:])
            i += 1
            continue

        add_runs(doc.add_paragraph(), stripped)
        i += 1
    return n_head, n_table


def build(src, out):
    with open(src, encoding="utf-8") as f:
        lines = f.readlines()
    doc = Document()
    nh, nt = render(lines, doc)
    doc.save(out)
    return nh, nt


def selftest():
    sample = [
        "# Title\n", "\n", "## Section\n", "intro **bold** text\n",
        "- item one\n", "  - nested **b**\n", "\n",
        "| a | b |\n", "|---|---|\n", "| 1 | 2 |\n", "\n", "### Sub\n",
    ]
    doc = Document()
    nh, nt = render(sample, doc)
    assert nh == 3, f"headings={nh}"
    assert nt == 1, f"tables={nt}"
    t = doc.tables[0]
    assert len(t.rows) == 2 and len(t.columns) == 2, "table shape"
    assert t.rows[0].cells[0].paragraphs[0].runs[0].bold, "header bold missing"
    assert not t.rows[1].cells[0].paragraphs[0].runs[0].bold, "data cell should not be bold"
    print(f"selftest OK: headings={nh} tables={nt}")


if __name__ == "__main__":
    if "--selftest" in sys.argv:
        selftest()
    else:
        nh, nt = build(SRC, OUT)
        print(f"wrote {OUT} (headings={nh} tables={nt})")
